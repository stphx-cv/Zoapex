#!/usr/bin/env python3
"""Genera el documento Word de entrega PA4 (Parte 4) para Zoapex."""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE = Path(__file__).resolve().parent
PA3_SRC = BASE.parent / "pa3" / "send" / "Zoapex-P1.docx"
OUT_DIR = BASE / "send"
OUT_FILE = OUT_DIR / "Zoapex-PA3-PA4.docx"

COLOR_TITLE = RGBColor(0x55, 0x55, 0x55)
COLOR_HEADING = RGBColor(0x2F, 0x5D, 0x8A)
COLOR_BODY = RGBColor(0x22, 0x22, 0x22)
COLOR_CODE_BG = RGBColor(0xF4, 0xF6, 0xF8)
FONT = "Arial"


def set_run(run, *, size=11, bold=False, color=COLOR_BODY, italic=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_paragraph(doc, text="", *, size=11, bold=False, color=COLOR_BODY, space_after=6, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 15, 2: 13, 3: 12}
    return add_paragraph(doc, text, size=sizes.get(level, 12), bold=True, color=COLOR_HEADING, space_after=8)


def add_body(doc, text, space_after=6):
    return add_paragraph(doc, text, size=11, color=COLOR_BODY, space_after=space_after)


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run("• " + text)
    set_run(run, size=11, color=COLOR_BODY)
    return p


def add_code_block(doc, code: str):
    for line in code.strip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    add_paragraph(doc, "", space_after=8)


def add_caption(doc, text):
    p = add_paragraph(doc, text, size=10, color=RGBColor(0x66, 0x66, 0x66), space_after=10)
    p.runs[0].font.italic = True
    return p


def add_placeholder(doc, text):
    p = add_paragraph(doc, text, size=10, bold=True, color=RGBColor(0xE7, 0x6F, 0x51), space_after=12)
    return p


def build_parte4(doc: Document):
    doc.add_page_break()

    add_heading(doc, "Parte 4 — Aplicación de la Sesión 12", level=1)
    add_body(
        doc,
        "Esta sección documenta la implementación de Entity Framework Core, LINQ, procedimientos "
        "almacenados (funciones PostgreSQL) y AJAX en la parte web del proyecto Zoapex, "
        "continuando el documento de planeamiento entregado en la PA3.",
    )

    add_heading(doc, "Decisión y justificación general", level=2)
    add_body(
        doc,
        "Para Zoapex mantuve la Ruta A (Escritorio + Web) definida en la PA3. El CRUD de productos, "
        "categorías y la operación maestro-detalle transaccional (Pedido → DetallePedido) se gestionan "
        "desde la aplicación de escritorio WPF con arquitectura en N capas y ADO.NET (Npgsql). "
        "La parte web (ASP.NET Core Razor Pages) se conecta a la misma base de datos en Supabase y "
        "permite al cliente ver el catálogo, iniciar sesión, armar un carrito y confirmar su compra.",
    )
    add_body(
        doc,
        "En la web utilizo EF Core como ORM principal para mapear las tablas de PostgreSQL, LINQ "
        "para las consultas del catálogo e historial de pedidos del cliente, y la función "
        "fn_register_order como procedimiento almacenado para registrar la venta de forma transaccional "
        "(cabecera + detalle + descuento de stock). El AJAX lo apliqué en el catálogo y en el carrito "
        "para filtrar productos, ver el detalle y actualizar totales sin recargar la página.",
    )
    add_body(
        doc,
        "Respeto la regla de oro: la vista nunca consulta la base de datos directamente. Las páginas "
        "Razor llaman a repositorios en la capa Data (CatalogRepository, OrderRepository, "
        "CustomerRepository), y estos usan EF Core, LINQ o la invocación al procedimiento almacenado.",
    )

    # A. EF Core
    add_heading(doc, "A. Entity Framework Core — DbContext y modelo de datos", level=2)
    add_body(
        doc,
        "Justificación: uso EF Core para mapear las tablas de Supabase (category, product, customer, "
        "order, order_detail) como entidades C# y centralizar las relaciones entre ellas. Esto me "
        "permite reutilizar el mismo contexto en todos los repositorios de la capa web.",
    )
    add_caption(doc, "Archivo: Zoapex.Web/Data/ZoapexDbContext.cs")
    add_code_block(
        doc,
        """public class ZoapexDbContext(DbContextOptions<ZoapexDbContext> options) : DbContext(options)
{
    public DbSet<Category> Categories => Set<Category>();
    public DbSet<Product> Products => Set<Product>();
    public DbSet<Customer> Customers => Set<Customer>();
    public DbSet<OrderHeader> Orders => Set<OrderHeader>();
    public DbSet<OrderDetail> OrderDetails => Set<OrderDetail>();

    protected override void OnModelCreating(ModelBuilder modelBuilder)
    {
        modelBuilder.Entity<Product>()
            .HasOne(p => p.Category)
            .WithMany(c => c.Products)
            .HasForeignKey(p => p.CategoryId);

        modelBuilder.Entity<OrderHeader>()
            .HasOne(o => o.Customer)
            .WithMany(c => c.Orders)
            .HasForeignKey(o => o.CustomerId);

        modelBuilder.Entity<OrderDetail>()
            .HasOne(d => d.Order)
            .WithMany(o => o.Details)
            .HasForeignKey(d => d.OrderId);
    }
}""",
    )
    add_placeholder(doc, "[CAPTURA: pantalla del archivo ZoapexDbContext.cs en Visual Studio / Cursor]")
    add_body(
        doc,
        "Explicación: el DbContext define los DbSet de cada tabla y configura las relaciones "
        "Product → Category, Order → Customer y OrderDetail → Order/Product. Las entidades usan "
        "atributos [Table] y [Column] para mapear los nombres en snake_case de PostgreSQL.",
    )

    # B. LINQ
    add_heading(doc, "B. LINQ — catálogo e historial de pedidos", level=2)
    add_body(
        doc,
        "Justificación: para listados y filtros frecuentes (catálogo por categoría, búsqueda por "
        "nombre/código y historial del cliente) uso LINQ to Entities porque la consulta queda legible "
        "en C# y el compilador valida los nombres de propiedades.",
    )
    add_caption(doc, "Archivo: Zoapex.Web/Data/CatalogRepository.cs — filtro del catálogo")
    add_code_block(
        doc,
        """public async Task<List<ProductCardDto>> GetCatalogAsync(string? search, int? categoryId)
{
    var query = ctx.Products
        .AsNoTracking()
        .Include(p => p.Category)
        .Where(p => p.Status == 1);

    if (categoryId is > 0)
        query = query.Where(p => p.CategoryId == categoryId);

    if (!string.IsNullOrWhiteSpace(search))
    {
        var term = search.Trim().ToLower();
        query = query.Where(p =>
            p.Name.ToLower().Contains(term) ||
            p.Code.ToLower().Contains(term) ||
            (p.Category != null && p.Category.Name.ToLower().Contains(term)));
    }

    return await query
        .OrderBy(p => p.Code)
        .Select(p => new ProductCardDto(...))
        .ToListAsync();
}""",
    )
    add_placeholder(doc, "[CAPTURA: método GetCatalogAsync en CatalogRepository.cs]")
    add_body(
        doc,
        "Explicación: esta consulta LINQ filtra productos activos por categoría y texto de búsqueda, "
        "hace un JOIN implícito con Category mediante Include y proyecta el resultado a un DTO para la vista.",
    )

    add_caption(doc, "Archivo: Zoapex.Web/Data/OrderRepository.cs — historial del cliente")
    add_code_block(
        doc,
        """public async Task<List<OrderHistoryDto>> GetCustomerOrdersAsync(int customerId, int take = 20)
{
    return await ctx.Orders
        .AsNoTracking()
        .Include(o => o.Details)
        .ThenInclude(d => d.Product)
        .Where(o => o.Status == 1 && o.CustomerId == customerId)
        .OrderByDescending(o => o.OrderDate)
        .Take(take)
        .Select(o => new OrderHistoryDto(...))
        .ToListAsync();
}""",
    )
    add_placeholder(doc, "[CAPTURA: método GetCustomerOrdersAsync en OrderRepository.cs]")
    add_body(
        doc,
        "Explicación: trae los pedidos del cliente autenticado con sus líneas de detalle y el nombre "
        "del producto, ordenados del más reciente al más antiguo. Se usa en la página Mis pedidos.",
    )

    # C. SP
    add_heading(doc, "C. Procedimiento almacenado — registro transaccional del pedido", level=2)
    add_body(
        doc,
        "Justificación: el registro de la venta (cabecera + detalle + cálculo de IGV + descuento de "
        "stock) involucra varias operaciones que deben ejecutarse 'todo o nada'. Por eso delego esa "
        "lógica a la función fn_register_order en PostgreSQL (equivalente a un stored procedure en "
        "SQL Server), reutilizando la misma función que ya usa la aplicación de escritorio WPF.",
    )
    add_caption(doc, "1) Función creada en Supabase (zoapex-db/sql/zoapex_database.sql)")
    add_code_block(
        doc,
        """CREATE OR REPLACE FUNCTION fn_register_order(
    p_customer_id INT,
    p_details     JSONB
)
RETURNS INT AS $$
BEGIN
    -- Calcula subtotal, IGV (18%) y total
    -- Inserta cabecera en "order"
    -- Inserta líneas en order_detail
    -- Descuenta stock de cada producto
    RETURN v_order_id;
END;
$$ LANGUAGE plpgsql;""",
    )
    add_placeholder(doc, "[CAPTURA: script fn_register_order en Supabase SQL Editor o en el archivo .sql]")
    add_caption(doc, "2) Invocación desde el proyecto web (EF Core + SqlQuery)")
    add_code_block(
        doc,
        """var json = JsonSerializer.Serialize(payload);

var orderId = await ctx.Database
    .SqlQuery<int>($"SELECT fn_register_order({customerId}, {json}::jsonb) AS \\"Value\\"")
    .SingleAsync();

return orderId;""",
    )
    add_placeholder(doc, "[CAPTURA: método RegisterOrderAsync en OrderRepository.cs]")
    add_body(
        doc,
        "Explicación: la capa web serializa las líneas del carrito a JSON, invoca fn_register_order "
        "mediante SqlQuery y recibe el ID del pedido registrado. Si algo falla en la base de datos, "
        "no se confirma la venta.",
    )

    # D. AJAX
    add_heading(doc, "D. AJAX (obligatorio) — catálogo y carrito sin recargar", level=2)
    add_body(
        doc,
        "Justificación: al navegar el catálogo el cliente necesita filtrar productos y ver precio/stock "
        "al instante. En el carrito, al quitar un ítem debe actualizarse el resumen de totales sin "
        "recargar toda la página. Por eso implementé handlers Razor Pages que devuelven JSON y fetch "
        "en JavaScript.",
    )
    add_caption(doc, "1) Handler AJAX en el code-behind (Catalog.cshtml.cs)")
    add_code_block(
        doc,
        """public async Task<IActionResult> OnGetSearchAsync(string? q, int? categoryId)
{
    var products = await catalogRepo.GetCatalogAsync(q, categoryId);
    return new JsonResult(products);
}

public async Task<IActionResult> OnGetProductAsync(int id)
{
    var product = await catalogRepo.GetProductDetailAsync(id);
    if (product is null)
        return new JsonResult(new { found = false });

    return new JsonResult(new { found = true, product.Name, product.Price, product.Stock, ... });
}""",
    )
    add_placeholder(doc, "[CAPTURA: handlers OnGetSearchAsync y OnGetProductAsync]")
    add_caption(doc, "2) Fetch en la vista (Catalog.cshtml)")
    add_code_block(
        doc,
        """async function loadProducts() {
    const url = `?handler=Search&q=${q}&categoryId=${categoryId}`;
    const response = await fetch(url);
    const products = await response.json();
    // actualiza solo la grilla de productos, sin recargar la página
    grid.innerHTML = products.map(p => `...`).join('');
}

async function showProductDetail(id) {
    const response = await fetch(`?handler=Product&id=${id}`);
    const data = await response.json();
    // actualiza nombre, precio y stock en el panel lateral
}""",
    )
    add_placeholder(doc, "[CAPTURA: funciones JavaScript loadProducts y showProductDetail en Catalog.cshtml]")
    add_body(
        doc,
        "Explicación: al escribir en el buscador o cambiar la categoría, fetch llama al handler "
        "OnGetSearchAsync y actualiza solo la grilla de productos. Al pulsar 'Ver detalle', "
        "OnGetProductAsync devuelve el precio y stock sin recargar la página. En el carrito, "
        "refreshCart() usa ?handler=Summary para recalcular subtotal, IGV y total vía AJAX.",
    )
    add_placeholder(doc, "[CAPTURA: pantalla del catálogo filtrando productos en el navegador]")
    add_placeholder(doc, "[CAPTURA: pantalla del carrito con productos y totales]")
    add_placeholder(doc, "[CAPTURA: pantalla de login del cliente y confirmación de pedido]")

    # Login cliente
    add_heading(doc, "E. Login de cliente (fase web)", level=2)
    add_body(
        doc,
        "Complemento planificado en PA3: la web incluye registro e inicio de sesión de clientes. "
        "CustomerRepository valida credenciales con LINQ, PasswordHelper encripta la contraseña y "
        "CustomerSession guarda la sesión. El checkout del carrito requiere login y asocia el pedido "
        "al customer_id del cliente autenticado.",
    )
    add_placeholder(doc, "[CAPTURA: página Account/Login.cshtml y Mis pedidos del cliente]")

    # Checklist
    add_heading(doc, "Checklist de requisitos mínimos (PA4)", level=2)
    add_bullet(doc, "La parte web accede a datos con EF Core, LINQ y procedimiento almacenado (fn_register_order).")
    add_bullet(doc, "Al menos una funcionalidad usa AJAX (catálogo: filtro y detalle; carrito: resumen).")
    add_bullet(doc, "Cada tecnología incluye justificación escrita y evidencia de código.")
    add_bullet(doc, "La UI nunca consulta la base de datos directamente; siempre pasa por repositorios en Data/.")
    add_bullet(doc, "Proyecto comprimido en .zip adjunto a la entrega.")


def update_header(doc: Document):
    """Actualiza nombre y fecha en el encabezado si existe placeholder."""
    for p in doc.paragraphs:
        if p.text.startswith("Alumno:"):
            p.clear()
            r = p.add_run("Alumno:  Stephano Camarena Villa\tFecha de entrega:  2 de julio de 2026")
            set_run(r, size=11, color=COLOR_BODY)
        if p.text.startswith("Fecha:") and "Proyecto:" in p.text:
            pass


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    if PA3_SRC.exists():
        doc = Document(str(PA3_SRC))
        update_header(doc)
    else:
        doc = Document()
        add_paragraph(doc, "TAREA — TRABAJO DE PLANEAMIENTO", size=10, bold=True, color=COLOR_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_paragraph(doc, "Zoapex — Documento de planeamiento", size=20, bold=True, color=COLOR_HEADING, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_paragraph(doc, "Desarrollo de Aplicaciones Básico", size=11, color=COLOR_BODY, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_paragraph(doc, "Alumno:  Stephano Camarena Villa\tFecha de entrega:  2 de julio de 2026", size=11, color=COLOR_BODY)

    build_parte4(doc)
    doc.save(str(OUT_FILE))
    print(f"Generado: {OUT_FILE}")


if __name__ == "__main__":
    main()
