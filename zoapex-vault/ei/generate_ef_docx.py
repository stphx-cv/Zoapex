#!/usr/bin/env python3
"""Genera el documento Word del Examen Final para Zoapex.

Parte del documento de la PA4 (que ya contiene PA3 + PA4) y le agrega una
sección más: "Examen Final — Transaccionalidad Web + Reporte a Excel".
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE = Path(__file__).resolve().parent
PA4_SRC = BASE.parent / "pa4" / "send" / "Zoapex-PA4.docx"
OUT_DIR = BASE / "send"
OUT_FILE = OUT_DIR / "Zoapex-ExamenFinal.docx"

COLOR_TITLE = RGBColor(0x55, 0x55, 0x55)
COLOR_HEADING = RGBColor(0x2F, 0x5D, 0x8A)
COLOR_BODY = RGBColor(0x22, 0x22, 0x22)
FONT = "Arial"


def set_run(run, *, size=11, bold=False, color=COLOR_BODY, italic=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_paragraph(doc, text="", *, size=11, bold=False, color=COLOR_BODY, space_after=6, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 15, 2: 13, 3: 12}
    return add_paragraph(doc, text, size=sizes.get(level, 12), bold=True, color=COLOR_HEADING, space_after=8)


def add_body(doc, text, space_after=6):
    return add_paragraph(doc, text, size=11, color=COLOR_BODY, space_after=space_after)


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run("• " + text)
    set_run(run, size=11, color=COLOR_BODY)
    return p


def add_code_block(doc, code: str):
    for line in code.strip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    add_paragraph(doc, "", space_after=8)


def add_caption(doc, text):
    p = add_paragraph(doc, text, size=10, color=RGBColor(0x66, 0x66, 0x66), space_after=10)
    p.runs[0].font.italic = True
    return p


def add_placeholder(doc, text):
    return add_paragraph(doc, text, size=10, bold=True, color=RGBColor(0xE7, 0x6F, 0x51), space_after=12)


def build_examen_final(doc: Document):
    doc.add_page_break()

    # Portada de la sección (nombre, proyecto y fecha)
    add_paragraph(doc, "EVALUACIÓN FINAL", size=10, bold=True, color=COLOR_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Cierra tu proyecto: Transaccionalidad Web + Reporte a Excel",
                  size=18, bold=True, color=COLOR_HEADING, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Desarrollo de Aplicaciones Básico", size=11, color=COLOR_BODY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Alumno:  Stephano Camarena Villa", size=11, color=COLOR_BODY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Proyecto:  Zoapex — Tienda de mascotas", size=11, color=COLOR_BODY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Fecha de entrega:  16 de julio de 2026", size=11, color=COLOR_BODY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "", space_after=10)

    add_body(
        doc,
        "Esta sección cierra el ciclo del dato del proyecto Zoapex aplicando el flujo guardar → "
        "mostrar → exportar sobre la misma solución de la PA3/PA4. Se implementa una operación "
        "transaccional maestro-detalle en la web, un reporte de ventas con análisis y la descarga de "
        "ese reporte a Excel con EPPlus. Se respeta la regla de oro: la vista nunca consulta la base "
        "de datos directamente, siempre pasa por la capa de datos (repositorios y el servicio de "
        "exportación).",
    )

    # ---- Parte A ----
    add_heading(doc, "A. Escenario transaccional en la web (maestro-detalle, 'todo o nada')", level=2)
    add_body(
        doc,
        "Qué se guarda: una Venta / Pedido con estructura cabecera-detalle: la tabla order (cabecera) "
        "y order_detail (varias líneas). La operación se graba 'todo o nada': si una línea falla, no "
        "se inserta nada. La cabecera, todos los detalles y el descuento de stock se ejecutan dentro "
        "de la misma transacción del procedimiento almacenado en PostgreSQL (fn_register_order).",
    )
    add_caption(doc, "1) La vista llama a la capa de datos, nunca al DbContext (Pages/Cart.cshtml.cs)")
    add_code_block(
        doc,
        """var orderId = await orderRepo.RegisterOrderAsync(customerId, lines);
CartSession.Clear(HttpContext.Session);
TempData["Success"] = $"¡Pedido registrado correctamente! Número de orden: {orderId}";
return RedirectToPage("/Orders");""",
    )
    add_caption(doc, "2) El repositorio delega la escritura transaccional al SP (Data/OrderRepository.cs)")
    add_code_block(
        doc,
        """public async Task<int> RegisterOrderAsync(int? customerId, IReadOnlyList<CartLineDto> lines)
{
    // ...validaciones de cantidad y stock...
    var payload = lines.Select(l => new {
        product_id = l.ProductId, quantity = l.Quantity, unit_price = l.UnitPrice });
    var json = JsonSerializer.Serialize(payload);

    var orderId = await ctx.Database
        .SqlQuery<int>($"SELECT fn_register_order({customerId}, {json}::jsonb) AS \\"Value\\"")
        .SingleAsync();
    return orderId;
}""",
    )
    add_caption(doc, "3) Transacción 'todo o nada' en PostgreSQL (zoapex-db/sql/zoapex_database.sql)")
    add_code_block(
        doc,
        """CREATE OR REPLACE FUNCTION fn_register_order(p_customer_id INT, p_details JSONB)
RETURNS INT AS $$
DECLARE v_order_id INT; ...
BEGIN
    -- Genera código y calcula subtotal / IGV / total recorriendo el JSON
    INSERT INTO "order" (code, customer_id, subtotal, tax, total, status)
    VALUES (v_code, p_customer_id, v_subtotal, v_tax, v_total, 1)
    RETURNING order_id INTO v_order_id;

    -- Inserta detalles y descuenta stock (misma transacción)
    FOR v_detail IN SELECT * FROM jsonb_array_elements(p_details) LOOP
        INSERT INTO order_detail (order_id, product_id, quantity, unit_price, subtotal)
        VALUES (v_order_id, v_prod_id, v_qty, v_price, v_line_subtotal);
        UPDATE product SET stock = stock - v_qty WHERE product_id = v_prod_id;
    END LOOP;

    RETURN v_order_id;
END;
$$ LANGUAGE plpgsql;""",
    )
    add_placeholder(doc, "[CAPTURA A-1: carrito y mensaje '¡Pedido registrado correctamente!' en Mis pedidos]")
    add_placeholder(doc, "[CAPTURA A-2: registros en las tablas order y order_detail con el stock descontado]")

    # ---- Parte B ----
    add_heading(doc, "B. Reporte y descarga a Excel con EPPlus", level=2)

    add_heading(doc, "B.1. El reporte en la web (mostrar)", level=3)
    add_body(
        doc,
        "Cree una página nueva /SalesReport ('Reporte de ventas') que presenta el dato como análisis, "
        "tabla y gráfico (no como texto crudo). Muestra cuatro indicadores (KPIs): número de pedidos, "
        "unidades vendidas, ticket promedio y ventas totales; un gráfico de barras con el Top 5 de "
        "productos por ingreso; una tabla de ventas por fecha; y una tabla con el detalle de pedidos.",
    )

    add_heading(doc, "B.2. Cómo llegué al reporte (explicado con mis palabras)", level=3)
    add_body(
        doc,
        "Partí del dato que ya guardo en la Parte A: los pedidos (order) con sus líneas (order_detail). "
        "Me pregunté qué necesita un administrador para entender el negocio y decidí no mostrar los "
        "pedidos 'en crudo', sino resumirlos y analizarlos. Primero armé una sola consulta LINQ "
        "(GetSalesReportAsync) que trae los pedidos activos; a partir de esa información calculé los "
        "cuatro KPIs, siendo el ticket promedio las ventas totales divididas entre el número de "
        "pedidos. Luego quise responder '¿qué se vende más?', así que agrupé las líneas por producto y "
        "sumé sus ingresos para obtener el Top 5, que muestro como gráfico de barras (cada barra es "
        "proporcional al ingreso del producto respecto al mayor). Para ver la evolución en el tiempo, "
        "agrupé los mismos pedidos por fecha y sumé sus totales. Finalmente dejé una tabla con el "
        "detalle de pedidos para revisar caso por caso. Así el reporte pasa de 'texto crudo' a un "
        "tablero con análisis, gráfico y tablas, y ese mismo dato es el que después se descarga en Excel.",
    )
    add_placeholder(doc, "[CAPTURA B-1: página /SalesReport con KPIs, gráfico de barras y tablas]")

    add_heading(doc, "B.3. La descarga a Excel con EPPlus (exportar)", level=3)
    add_body(
        doc,
        "El botón 'Descargar Excel' apunta al handler OnGetExportAsync, que genera y descarga un "
        "archivo .xlsx real. La respuesta usa el content-type de Office y un nombre de archivo con "
        "fecha y hora.",
    )
    add_caption(doc, "1) Handler que devuelve el archivo (Pages/SalesReport.cshtml.cs)")
    add_code_block(
        doc,
        """public async Task<IActionResult> OnGetExportAsync()
{
    if (!CustomerSession.IsAdmin(HttpContext.Session))
        return RedirectToPage("/Account/Login", new { returnUrl = "/SalesReport" });

    var report = await orderRepo.GetSalesReportAsync();
    var bytes  = exporter.Build(report);
    return File(bytes,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        $"reporte-ventas-{DateTime.Now:yyyyMMdd-HHmm}.xlsx");
}""",
    )
    add_caption(doc, "2) Armado del libro con EPPlus (Data/SalesExcelExporter.cs, extracto)")
    add_code_block(
        doc,
        """public byte[] Build(SalesReportDto data)
{
    using var package = new ExcelPackage();
    var ws = package.Workbook.Worksheets.Add("Resumen");
    ws.Cells["A1"].Value = "Zoapex — Reporte de ventas";
    // ...KPIs con formato moneda "S/ #,##0.00"...
    var chart = ws.Drawings.AddBarChart("TopProductsChart", eBarChartType.BarClustered);
    chart.Series.Add(/* ingresos */, /* nombres de producto */);
    // Hoja "Detalle de pedidos" con AutoFilter + fila TOTAL con SUM(...)
    return package.GetAsByteArray();
}""",
    )
    add_caption(doc, "3) Licencia (uso académico / no comercial) y registro del servicio (Program.cs)")
    add_code_block(
        doc,
        """ExcelPackage.License.SetNonCommercialPersonal("Stephano Camarena Villa");
builder.Services.AddScoped<SalesExcelExporter>();""",
    )
    add_placeholder(doc, "[CAPTURA B-2: código del SalesExcelExporter.cs en el editor]")
    add_placeholder(doc, "[CAPTURA B-3: Excel descargado abierto, hojas 'Resumen' y 'Detalle de pedidos']")

    add_heading(doc, "B.4. La query adjunta (LINQ que alimenta el reporte)", level=3)
    add_body(
        doc,
        "Consulta LINQ de agregación en la capa de datos (Data/OrderRepository.cs, método "
        "GetSalesReportAsync). Trae los pedidos con EF Core y calcula los KPIs, el top de productos y "
        "las ventas por fecha.",
    )
    add_code_block(
        doc,
        """// Filas de pedidos (cabecera + nombre de cliente + n.º de ítems)
var orders = await ctx.Orders.AsNoTracking()
    .Where(o => o.Status == 1)
    .OrderByDescending(o => o.OrderDate)
    .Select(o => new SalesOrderRowDto(
        o.Code, o.OrderDate,
        o.Customer != null ? o.Customer.FirstName + " " + o.Customer.LastName : "Mostrador",
        o.Details.Sum(d => (int?)d.Quantity) ?? 0,
        o.Subtotal, o.Tax, o.Total))
    .ToListAsync();

// Detalle plano + Top 5 productos por ingreso
var detailRows = await ctx.OrderDetails.AsNoTracking()
    .Where(d => d.Order.Status == 1)
    .Select(d => new { d.Product.Name, d.Quantity, d.Subtotal })
    .ToListAsync();
var topProducts = detailRows
    .GroupBy(d => d.Name)
    .Select(g => new TopProductDto(g.Key, g.Sum(d => d.Quantity), g.Sum(d => d.Subtotal)))
    .OrderByDescending(t => t.Revenue).Take(5).ToList();""",
    )
    add_caption(doc, "Equivalente en SQL de la agregación principal (Top productos)")
    add_code_block(
        doc,
        """SELECT p.name AS producto,
       SUM(d.quantity) AS cantidad,
       SUM(d.subtotal) AS ingreso
FROM   order_detail d
JOIN   "order" o ON o.order_id = d.order_id
JOIN   product p ON p.product_id = d.product_id
WHERE  o.status = 1
GROUP  BY p.name
ORDER  BY ingreso DESC
LIMIT  5;""",
    )

    # ---- Parte C ----
    add_heading(doc, "C. Seguridad (opcional — puntos extra)", level=2)
    add_bullet(doc, "Roles Admin / Cliente: columna 'role' en customer; el rol se guarda en la sesión al iniciar sesión.")
    add_bullet(doc, "Reglas de acceso: la página y el handler de exportación solo los ve el Administrador (guardia CustomerSession.IsAdmin); el enlace 'Reporte' del menú solo aparece para Admin.")
    add_bullet(doc, "Formulario de Login (/Account/Login) con validación de credenciales y hash de contraseña (PasswordHelper).")
    add_bullet(doc, "Caducidad de sesión (timeout): IdleTimeout de 30 minutos configurado en Program.cs.")
    add_body(doc, "Cuentas demo — Admin: admin@zoapex.com / zoapex123 · Cliente: cliente@zoapex.com / zoapex123.", space_after=8)
    add_placeholder(doc, "[CAPTURA C-1: menú 'Reporte' visible como Admin; acceso denegado (redirección a Login) como Cliente]")

    # ---- Verificación ----
    add_heading(doc, "Verificación realizada (de punta a punta)", level=2)
    add_bullet(doc, "dotnet build → compila sin errores.")
    add_bullet(doc, "Login como admin@zoapex.com → /SalesReport responde HTTP 200 con KPIs, gráfico y tablas.")
    add_bullet(doc, "Botón Descargar Excel → descarga un .xlsx real (hojas 'Resumen' y 'Detalle de pedidos', con gráfico y autofiltro).")
    add_bullet(doc, "Acceso sin sesión a /SalesReport y al handler Export → HTTP 302 hacia el Login (regla de acceso por rol funcionando).")

    # ---- Entregable ----
    add_heading(doc, "Entregable", level=2)
    add_bullet(doc, "Proyecto actualizado en .zip (carpeta zoapex-web/), o video corto: guardar carrito → ver reporte → descargar Excel.")
    add_bullet(doc, "Este documento (con capturas) adjunto al de la PA3/PA4.")
    add_bullet(doc, "Excel de muestra generado por el propio sistema: reporte-ventas-ejemplo.xlsx.")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    if PA4_SRC.exists():
        doc = Document(str(PA4_SRC))
    else:
        doc = Document()
        add_paragraph(doc, "Zoapex — Documento del proyecto", size=20, bold=True,
                      color=COLOR_HEADING, align=WD_ALIGN_PARAGRAPH.CENTER)

    build_examen_final(doc)
    doc.save(str(OUT_FILE))
    print(f"Generado: {OUT_FILE}")


if __name__ == "__main__":
    main()
